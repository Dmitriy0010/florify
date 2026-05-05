import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
MY_DOC_PATH = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\5-текст.docx"
OUTPUT_PATH = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\5-текст_updated.docx"
SCHEMA_PATH = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\_docx_tools\table_schema.json"

# Load schema
with open(SCHEMA_PATH, 'r', encoding='utf-8') as f:
    schema = json.load(f)

# Table mapping and descriptions (consolidated)
table_order = [
    "product_categories", "products", "users", "user_roles", "refresh_tokens",
    "stores", "customers", "loyalty_accounts", "loyalty_transactions",
    "orders", "order_items", "stock_balances", "stock_batches", "stock_transactions",
    "suppliers", "purchase_invoices", "purchase_invoice_items", "employees",
    "employee_salary_configs", "employee_timesheet", "employee_salary_statements",
    "delivery_zones", "delivery_slots", "delivery_tasks", "financial_transactions",
    "analytics_order_facts", "analytics_cost_facts", "media_files",
    "notification_templates", "notification_logs", "recipes", "recipe_items",
    "payments", "shedlock"
]

table_titles = {
    "product_categories": "Категории товаров", "products": "Товары", "users": "Пользователи",
    "user_roles": "Роли пользователей", "refresh_tokens": "Токены обновления", "stores": "Магазины",
    "customers": "Клиенты", "loyalty_accounts": "Аккаунты лояльности", "loyalty_transactions": "Транзакции лояльности",
    "orders": "Заказы", "order_items": "Позиции заказов", "stock_balances": "Остатки на складе",
    "stock_batches": "Партии товаров", "stock_transactions": "Складские транзакции", "suppliers": "Поставщики",
    "purchase_invoices": "Приходные накладные", "purchase_invoice_items": "Позиции накладных",
    "employees": "Сотрудники", "employee_salary_configs": "Конфигурации зарплат", "employee_timesheet": "Табель учета времени",
    "employee_salary_statements": "Зарплатные ведомости", "delivery_zones": "Зоны доставки",
    "delivery_slots": "Слоты доставки", "delivery_tasks": "Задания на доставку", "financial_transactions": "Финансовые транзакции",
    "analytics_order_facts": "Факты заказов (аналитика)", "analytics_cost_facts": "Факты затрат (аналитика)",
    "media_files": "Медиафайлы", "notification_templates": "Шаблоны уведомлений", "notification_logs": "Логи уведомлений",
    "recipes": "Рецепты", "recipe_items": "Позиции рецептов", "payments": "Платежи", "shedlock": "Распределенные блокировки"
}

table_descriptions = {
    "product_categories": "Справочник категорий товаров. Позволяет структурировать каталог и упрощает поиск.",
    "products": "Центральный каталог товаров — каждый цветок, букет или аксессуар. Содержит актуальные цены и характеристики.",
    "users": "Аккаунты всех пользователей системы. Обеспечивает безопасность и разграничение доступа.",
    "user_roles": "Реализация ролевой модели доступа (RBAC). Позволяет назначать права пользователям.",
    "refresh_tokens": "Хранение токенов обновления для JWT-аутентификации. Повышает безопасность сессий.",
    "stores": "Реестр филиалов магазина. Является корневой сущностью для большинства операционных таблиц.",
    "customers": "CRM-профили клиентов. Позволяет хранить историю взаимодействий и маркетинговые теги.",
    "loyalty_accounts": "Счета программы лояльности. Отслеживает баланс баллов и уровень клиента.",
    "loyalty_transactions": "История всех операций по бонусным счетам. Обеспечивает прозрачность программы лояльности.",
    "orders": "Центральная сущность заказов. Хранит информацию о статусе, оплате и способе доставки.",
    "order_items": "Позиции заказов. Фиксирует состав, цены и количество на момент покупки.",
    "stock_balances": "Актуальные остатки товаров на складах. Использует метод средневзвешенной себестоимости (WAC).",
    "stock_batches": "Партии товаров. Позволяет отслеживать сроки годности и автоматически списывать просрочку.",
    "stock_transactions": "Журнал всех движений товаров. Основа для инвентаризаций и аудита.",
    "suppliers": "Справочник поставщиков. Содержит контактные данные и рейтинг надежности.",
    "purchase_invoices": "Входящие накладные. Основание для оприходования товара на склад.",
    "purchase_invoice_items": "Детализация накладных по товарам с указанием заказанного и принятого количества.",
    "employees": "Профили сотрудников. Содержит информацию о должности и привязке к филиалу.",
    "employee_salary_configs": "Настройки расчета зарплаты. Поддерживает оклад, процент и бонусы.",
    "employee_timesheet": "Табель учета времени. Фиксирует смены и отработанные часы.",
    "employee_salary_statements": "Зарплатные ведомости. Расчет финальных сумм на основе табеля.",
    "delivery_zones": "Географические зоны обслуживания с указанием стоимости доставки.",
    "delivery_slots": "Временные интервалы доставки. Балансирует нагрузку на курьеров.",
    "delivery_tasks": "Задания на доставку. Содержит статус, координаты и время прибытия.",
    "financial_transactions": "Журнал финансовых операций. Регистрирует выручку и расходы.",
    "analytics_order_facts": "Данные о заказах для быстрой аналитики выручки и маржинальности.",
    "analytics_cost_facts": "Данные о затратах (закупки, зарплаты) для анализа структуры расходов.",
    "media_files": "Метаданные загруженных медиа-файлов (фотографии товаров).",
    "notification_templates": "Шаблоны уведомлений (SMS, Email, Push). Поддерживает переменные.",
    "notification_logs": "Журнал отправленных уведомлений. Отслеживает статус доставки сообщений.",
    "recipes": "Технологические карты (состав букетов). Позволяет списывать ингредиенты.",
    "recipe_items": "Составляющие рецепта с указанием необходимого количества ингредиентов.",
    "payments": "Транзакции платежных систем. Хранит данные об интеграции и статусы оплаты.",
    "shedlock": "Распределенные блокировки для фоновых задач в микросервисной среде."
}

column_desc_map = {
    "id": "уникальный идентификатор", "name": "наименование", "description": "описание",
    "created_at": "дата создания", "updated_at": "дата обновления", "active": "активность",
    "sku": "артикул", "category_id": "ID категории", "unit": "ед. изм.", "current_price": "цена",
    "image_url": "URL фото", "email": "email", "phone": "телефон", "password_hash": "хэш пароля",
    "role": "роль", "address": "адрес", "status": "статус", "total_amount": "сумма",
    "customer_id": "ID клиента", "store_id": "ID магазина", "product_id": "ID товара",
    "quantity": "количество", "unit_price": "цена за ед.", "points_balance": "баллы",
    "expires_at": "срок годности", "order_id": "ID заказа", "type": "тип", "occurred_at": "дата события",
    "amount": "сумма", "tax_id": "ИНН", "rating": "рейтинг", "hire_date": "дата приема",
    "dismiss_date": "дата увольнения", "base_amount": "оклад", "sales_percent": "% продаж",
    "bonus_per_order": "бонус", "date": "дата", "hours_worked": "часы", "delivery_fee": "цена доставки",
    "polygon": "координаты", "start_time": "начало", "end_time": "конец", "max_capacity": "лимит",
    "current_load": "нагрузка", "delivery_address": "адрес доставки", "latitude": "широта",
    "longitude": "долгота", "performer_id": "исполнитель", "write_off_reason": "причина списания",
    "supplier_id": "ID поставщика", "invoice_number": "номер накладной", "order_number": "номер заказа",
    "idempotency_key": "ключ", "is_paid": "оплачено", "is_active": "активно",
    "mime_type": "тип файла", "bucket": "бакет", "base_path": "путь", "uploaded_at": "дата загрузки",
    "code": "код", "channel": "канал", "subject": "тема", "body_template": "шаблон",
    "recipient_id": "ID получателя", "error_message": "ошибка", "qr_code_data": "QR-код",
    "confirmation_url": "URL оплаты", "external_id": "внешний ID"
}

def get_col_desc(col_name):
    return column_desc_map.get(col_name, f"поле {col_name}")

# Generate content
doc = Document(MY_DOC_PATH)

# Find target location
target_idx = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("3.1"):
        target_idx = i
        break

if target_idx == -1:
    target_idx = len(doc.paragraphs) - 1
else:
    # We want to insert after the introductory sentence (if any) or after the heading
    # Looking at the check_paras.py output, paragraph 93 is an intro. 
    # Let's insert after paragraph 93.
    target_idx = 93 

# Insert tables
for i, tname in enumerate(table_order, 1):
    title = table_titles[tname]
    desc = table_descriptions[tname]
    cols = schema.get(tname, [])
    attrs = "; ".join([f"{c['name']} – {get_col_desc(c['name'])}" for c in cols])
    
    text = f"{i}. {title} ({tname}). {desc} Содержит следующие атрибуты: {attrs}."
    
    # Insert
    if target_idx + 1 < len(doc.paragraphs):
        new_p = doc.paragraphs[target_idx + 1].insert_paragraph_before(text)
    else:
        new_p = doc.add_paragraph(text)
    
    # Format
    new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = new_p.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    
    # Font
    run = new_p.runs[0] if new_p.runs else new_p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    target_idx += 1 # Advance insertion point

doc.save(OUTPUT_PATH)
print("Done!")

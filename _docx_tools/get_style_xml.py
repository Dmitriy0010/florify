"""
Get the full XML of key styles and show full normal paragraph XML from example doc.
"""
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

EXAMPLE_PATH = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\5-текст отчета от Введения до Заключ.docx"
MY_DOC_PATH  = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\5-текст.docx"

ex_doc = Document(EXAMPLE_PATH)
my_doc = Document(MY_DOC_PATH)

# Get Normal style XML from EXAMPLE
print("=== EXAMPLE Normal style XML ===")
for s in ex_doc.styles:
    if s.name == 'Normal':
        print(etree.tostring(s.element, pretty_print=True).decode('utf-8'))
        break

# Get Normal style XML from MY DOC
print("=== MY DOC Normal style XML ===")
for s in my_doc.styles:
    if s.name == 'Normal':
        print(etree.tostring(s.element, pretty_print=True).decode('utf-8'))
        break

# Show first real paragraph XML from EXAMPLE to understand default formatting
print("\n=== EXAMPLE - paragraph 34 XML (Normal paragraph) ===")
p = ex_doc.paragraphs[34]
print(etree.tostring(p._element, pretty_print=True).decode('utf-8'))

# Show MY DOC paragraph 93 XML (Normal paragraph near section 3.1)
print("\n=== MY DOC - paragraph 93 XML ===")
p = my_doc.paragraphs[93]
print("Text:", p.text[:50])
print(etree.tostring(p._element, pretty_print=True).decode('utf-8'))

# Show MY DOC paragraph 91 XML (Heading 2 - section 3)
print("\n=== MY DOC - paragraph 91 XML (Heading 2) ===")
p = my_doc.paragraphs[91]
print("Text:", p.text[:50])
print(etree.tostring(p._element, pretty_print=True).decode('utf-8'))

# Show MY DOC paragraph 92 XML (Heading 2 - section 3.1)
print("\n=== MY DOC - paragraph 92 XML (Heading 2 subsection) ===")
p = my_doc.paragraphs[92]
print("Text:", p.text[:50])
print(etree.tostring(p._element, pretty_print=True).decode('utf-8'))

# Show example - section around "Универсальная"
print("\n=== EXAMPLE - find section with table list ===")
for i, p in enumerate(ex_doc.paragraphs):
    if 'Универсальная' in p.text or 'Photos' in p.text:
        print(f"Found at [{i}]: '{p.text[:80]}'")
        print(etree.tostring(p._element, pretty_print=True).decode('utf-8'))
        print("--- NEXT ---")
        if i+1 < len(ex_doc.paragraphs):
            p2 = ex_doc.paragraphs[i+1]
            print(f"[{i+1}]: '{p2.text[:80]}'")
            print(etree.tostring(p2._element, pretty_print=True).decode('utf-8'))
        break

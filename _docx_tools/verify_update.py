from docx import Document

doc = Document(r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\5-текст_updated.docx")
for i in range(90, 110):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        print(f"[{i}] style='{p.style.name}' text='{p.text[:100]}'")

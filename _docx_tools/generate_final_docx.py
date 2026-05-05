import json
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

MY_DOC_PATH = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\5-текст.docx"
OUTPUT_PATH = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\5-текст_updated.docx"
SCHEMA_PATH = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\_docx_tools\table_schema.json"
SQL_BLOCKS_PATH = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\_docx_tools\table_sql_blocks.json"

with open(SCHEMA_PATH, 'r', encoding='utf-8') as f:
    schema = json.load(f)

with open(SQL_BLOCKS_PATH, 'r', encoding='utf-8') as f:
    sql_blocks = json.load(f)

# The same mappings as before
table_order = [
    "product_categories", "products", "users", "user_roles", "refresh_tokens",
    "stores", "customers", "loyalty_accounts", "loyalty_transactions",
    "orders", "order_items", "stock_balances", "stock_batches", "stock_transactions",
    "suppliers", "purchase_invoices", "purchase_invoice_items", "employees",
    "employee_salary_configs", "employee_timesheet", "employee_salary_statements",
    "delivery_zones", "delivery_slots", "delivery_tasks", "financial_transactions",
    "analytics_order_facts", "analytics_cost_facts", "media_files",
    "notification_templates", "notification_logs", "recipes", "recipe_items",
    "payments", "shedlock"
]

table_titles = {
    "product_categories": "Категории товаров", "products": "Товары", "users": "Пользователи",
    "user_roles": "Роли пользователей", "refresh_tokens": "Токены обновления", "stores": "Магазины",
    "customers": "Клиенты", "loyalty_accounts": "Аккаунты лояльности", "loyalty_transactions": "Транзакции лояльности",
    "orders": "Заказы", "order_items": "Позиции заказов", "stock_balances": "Остатки на складе",
    "stock_batches": "Партии товаров", "stock_transactions": "Складские транзакции", "suppliers": "Поставщики",
    "purchase_invoices": "Приходные накладные", "purchase_invoice_items": "Позиции накладных",
    "employees": "Сотрудники", "employee_salary_configs": "Конфигурации зарплат", "employee_timesheet": "Табель учета времени",
    "employee_salary_statements": "Зарплатные ведомости", "delivery_zones": "Зоны доставки",
    "delivery_slots": "Слоты доставки", "delivery_tasks": "Задания на доставку", "financial_transactions": "Финансовые транзакции",
    "analytics_order_facts": "Факты заказов (аналитика)", "analytics_cost_facts": "Факты затрат (аналитика)",
    "media_files": "Медиафайлы", "notification_templates": "Шаблоны уведомлений", "notification_logs": "Логи уведомлений",
    "recipes": "Рецепты", "recipe_items": "Позиции рецептов", "payments": "Платежи", "shedlock": "Распределенные блокировки"
}

table_descriptions = {
    "product_categories": "Справочник категорий товаров. Позволяет структурировать каталог и упрощает поиск.",
    "products": "Центральный каталог товаров — каждый цветок, букет или аксессуар. Содержит актуальные цены и характеристики.",
    "users": "Аккаунты всех пользователей системы. Обеспечивает безопасность и разграничение доступа.",
    "user_roles": "Реализация ролевой модели доступа (RBAC). Позволяет назначать права пользователям.",
    "refresh_tokens": "Хранение токенов обновления для JWT-аутентификации. Повышает безопасность сессий.",
    "stores": "Реестр филиалов магазина. Является корневой сущностью для большинства операционных таблиц.",
    "customers": "CRM-профили клиентов. Позволяет хранить историю взаимодействий и маркетинговые теги.",
    "loyalty_accounts": "Счета программы лояльности. Отслеживает баланс баллов и уровень клиента.",
    "loyalty_transactions": "История всех операций по бонусным счетам. Обеспечивает прозрачность программы лояльности.",
    "orders": "Центральная сущность заказов. Хранит информацию о статусе, оплате и способе доставки.",
    "order_items": "Позиции заказов. Фиксирует состав, цены и количество на момент покупки.",
    "stock_balances": "Актуальные остатки товаров на складах. Использует метод средневзвешенной себестоимости (WAC).",
    "stock_batches": "Партии товаров. Позволяет отслеживать сроки годности и автоматически списывать просрочку.",
    "stock_transactions": "Журнал всех движений товаров. Основа для инвентаризаций и аудита.",
    "suppliers": "Справочник поставщиков. Содержит контактные данные и рейтинг надежности.",
    "purchase_invoices": "Входящие накладные. Основание для оприходования товара на склад.",
    "purchase_invoice_items": "Детализация накладных по товарам с указанием заказанного и принятого количества.",
    "employees": "Профили сотрудников. Содержит информацию о должности и привязке к филиалу.",
    "employee_salary_configs": "Настройки расчета зарплаты. Поддерживает оклад, процент и бонусы.",
    "employee_timesheet": "Табель учета времени. Фиксирует смены и отработанные часы.",
    "employee_salary_statements": "Зарплатные ведомости. Расчет финальных сумм на основе табеля.",
    "delivery_zones": "Географические зоны обслуживания с указанием стоимости доставки.",
    "delivery_slots": "Временные интервалы доставки. Балансирует нагрузку на курьеров.",
    "delivery_tasks": "Задания на доставку. Содержит статус, координаты и время прибытия.",
    "financial_transactions": "Журнал финансовых операций. Регистрирует выручку и расходы.",
    "analytics_order_facts": "Данные о заказах для быстрой аналитики выручки и маржинальности.",
    "analytics_cost_facts": "Данные о затратах (закупки, зарплаты) для анализа структуры расходов.",
    "media_files": "Метаданные загруженных медиа-файлов (фотографии товаров).",
    "notification_templates": "Шаблоны уведомлений (SMS, Email, Push). Поддерживает переменные.",
    "notification_logs": "Журнал отправленных уведомлений. Отслеживает статус доставки сообщений.",
    "recipes": "Технологические карты (состав букетов). Позволяет списывать ингредиенты.",
    "recipe_items": "Составляющие рецепта с указанием необходимого количества ингредиентов.",
    "payments": "Транзакции платежных систем. Хранит данные об интеграции и статусы оплаты.",
    "shedlock": "Распределенные блокировки для фоновых задач в микросервисной среде."
}

column_desc_map = {
    "id": "уникальный идентификатор", "name": "наименование", "description": "описание",
    "created_at": "дата создания", "updated_at": "дата обновления", "active": "активность",
    "sku": "артикул", "category_id": "ID категории", "unit": "ед. изм.", "current_price": "цена",
    "image_url": "URL фото", "email": "email", "phone": "телефон", "password_hash": "хэш пароля",
    "role": "роль", "address": "адрес", "status": "статус", "total_amount": "сумма",
    "customer_id": "ID клиента", "store_id": "ID магазина", "product_id": "ID товара",
    "quantity": "количество", "unit_price": "цена за ед.", "points_balance": "баллы",
    "expires_at": "срок годности", "order_id": "ID заказа", "type": "тип", "occurred_at": "дата события",
    "amount": "сумма", "tax_id": "ИНН", "rating": "рейтинг", "hire_date": "дата приема",
    "dismiss_date": "дата увольнения", "base_amount": "оклад", "sales_percent": "% продаж",
    "bonus_per_order": "бонус", "date": "дата", "hours_worked": "часы", "delivery_fee": "цена доставки",
    "polygon": "координаты", "start_time": "начало", "end_time": "конец", "max_capacity": "лимит",
    "current_load": "нагрузка", "delivery_address": "адрес доставки", "latitude": "широта",
    "longitude": "долгота", "performer_id": "исполнитель", "write_off_reason": "причина списания",
    "supplier_id": "ID поставщика", "invoice_number": "номер накладной", "order_number": "номер заказа",
    "idempotency_key": "ключ", "is_paid": "оплачено", "is_active": "активно",
    "mime_type": "тип файла", "bucket": "бакет", "base_path": "путь", "uploaded_at": "дата загрузки",
    "code": "код", "channel": "канал", "subject": "тема", "body_template": "шаблон",
    "recipient_id": "ID получателя", "error_message": "ошибка", "qr_code_data": "QR-код",
    "confirmation_url": "URL оплаты", "external_id": "внешний ID"
}

def get_col_desc(col_name):
    return column_desc_map.get(col_name, f"поле {col_name}")

doc = Document(MY_DOC_PATH)

# Look for paragraph starting with 3.1
# but in real docs it might just be the 93rd paragraph.
# We'll just hardcode finding "3.1 Описание таблиц" or find the insertion point directly.
# Let's find "3.1 " inside the document.
target_idx = -1
for i, p in enumerate(doc.paragraphs):
    if "3.1" in p.text and "Описание" in p.text:
        target_idx = i
        break

if target_idx == -1:
    # try to fallback to our previous finding: paragraph 92 was the 3.1 header, 93 was the text
    # we'll insert after 93
    target_idx = 93

# Clean up previously inserted text if any
# (If we run this multiple times on the original document, there's no previously inserted text)
# Actually, the user gave us the original document again? No, we just use the original doc.

current_idx = target_idx + 1

for i, tname in enumerate(table_order, 1):
    title = table_titles.get(tname, tname)
    desc = table_descriptions.get(tname, "")
    cols = schema.get(tname, [])
    
    # Format attributes string
    attr_list = [f"{c['name']} – {get_col_desc(c['name'])}" for c in cols]
    if "created_at" not in [c['name'] for c in cols]:
        # sometimes migrations don't list it directly or my regex missed it
        pass
    attrs_str = "; ".join(attr_list)
    
    # 1. Main description paragraph
    desc_text = f"{i}. {title} ({tname}). {desc} Содержит следующие атрибуты: {attrs_str}."
    
    p_desc = doc.paragraphs[current_idx].insert_paragraph_before(desc_text)
    p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_desc.paragraph_format.first_line_indent = Cm(1.25)
    p_desc.paragraph_format.line_spacing = 1.5
    for run in p_desc.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
    current_idx += 1
    
    # 2. "SQL-конструкция для инициализации таблицы:"
    p_sql_intro = doc.paragraphs[current_idx].insert_paragraph_before("SQL-конструкция для инициализации таблицы:")
    p_sql_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_sql_intro.paragraph_format.first_line_indent = Cm(1.25)
    p_sql_intro.paragraph_format.line_spacing = 1.5
    for run in p_sql_intro.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
    current_idx += 1
    
    # 3. SQL block
    sql_text = sql_blocks.get(tname, f"CREATE TABLE {tname} (...);")
    p_sql = doc.paragraphs[current_idx].insert_paragraph_before(sql_text)
    p_sql.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sql.paragraph_format.first_line_indent = Cm(0)
    p_sql.paragraph_format.left_indent = Cm(1.25)
    p_sql.paragraph_format.line_spacing = 1.0
    for run in p_sql.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        
    current_idx += 1

doc.save(OUTPUT_PATH)
print("Updated document generated at:", OUTPUT_PATH)

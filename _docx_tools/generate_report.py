import json
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
MY_DOC_PATH = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\5-текст.docx"
OUTPUT_PATH = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\5-текст_updated.docx"
SCHEMA_PATH = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\_docx_tools\table_schema.json"

# Load schema
with open(SCHEMA_PATH, 'r', encoding='utf-8') as f:
    schema = json.load(f)

# Table name mapping (English to Russian Title)
table_titles = {
    "product_categories": "Категории товаров",
    "products": "Товары",
    "users": "Пользователи",
    "user_roles": "Роли пользователей",
    "refresh_tokens": "Токены обновления",
    "stores": "Магазины",
    "customers": "Клиенты",
    "loyalty_accounts": "Аккаунты лояльности",
    "loyalty_transactions": "Транзакции лояльности",
    "orders": "Заказы",
    "order_items": "Позиции заказов",
    "stock_balances": "Остатки на складе",
    "stock_batches": "Партии товаров",
    "stock_transactions": "Складские транзакции",
    "suppliers": "Поставщики",
    "purchase_invoices": "Приходные накладные",
    "purchase_invoice_items": "Позиции накладных",
    "employees": "Сотрудники",
    "employee_salary_configs": "Конфигурации зарплат",
    "employee_timesheet": "Табель учета времени",
    "employee_salary_statements": "Зарплатные ведомости",
    "delivery_zones": "Зоны доставки",
    "delivery_slots": "Слоты доставки",
    "delivery_tasks": "Задания на доставку",
    "financial_transactions": "Финансовые транзакции",
    "analytics_order_facts": "Факты заказов (аналитика)",
    "analytics_cost_facts": "Факты затрат (аналитика)",
    "media_files": "Медиафайлы",
    "notification_templates": "Шаблоны уведомлений",
    "notification_logs": "Логи уведомлений",
    "recipes": "Рецепты",
    "recipe_items": "Позиции рецептов",
    "payments": "Платежи",
    "shedlock": "Распределенные блокировки"
}

# Column name mapping patterns
column_desc_map = {
    "id": "уникальный идентификатор",
    "name": "наименование",
    "description": "описание",
    "created_at": "дата и время создания",
    "updated_at": "дата и время обновления",
    "active": "флаг активности",
    "sku": "артикул товара",
    "category_id": "идентификатор категории",
    "unit": "единица измерения",
    "current_price": "текущая цена",
    "image_url": "ссылка на изображение",
    "email": "адрес электронной почты",
    "phone": "номер телефона",
    "password_hash": "хэш пароля",
    "role": "роль пользователя",
    "address": "адрес",
    "status": "статус",
    "total_amount": "общая сумма",
    "customer_id": "идентификатор клиента",
    "store_id": "идентификатор магазина",
    "product_id": "идентификатор товара",
    "quantity": "количество",
    "unit_price": "цена за единицу",
    "points_balance": "баланс баллов",
    "expires_at": "дата истечения срока",
    "order_id": "идентификатор заказа",
    "type": "тип",
    "occurred_at": "дата и время события",
    "amount": "сумма",
    "tax_id": "ИНН (налоговый идентификатор)",
    "rating": "рейтинг",
    "hire_date": "дата приема на работу",
    "dismiss_date": "дата увольнения",
    "base_amount": "базовая сумма",
    "sales_percent": "процент от продаж",
    "bonus_per_order": "бонус за заказ",
    "date": "дата",
    "hours_worked": "отработанные часы",
    "delivery_fee": "стоимость доставки",
    "polygon": "географические координаты зоны",
    "start_time": "время начала",
    "end_time": "время окончания",
    "max_capacity": "максимальная вместимость",
    "current_load": "текущая загрузка",
    "delivery_address": "адрес доставки",
    "latitude": "широта",
    "longitude": "долгота",
    "performer_id": "идентификатор исполнителя",
    "write_off_reason": "причина списания",
    "supplier_id": "идентификатор поставщика",
    "invoice_number": "номер накладной",
    "order_number": "номер заказа",
    "idempotency_key": "ключ идемпотентности",
    "is_paid": "флаг оплаты",
    "is_active": "признак активности",
    "mime_type": "MIME-тип файла",
    "bucket": "хранилище (баккет)",
    "base_path": "путь к файлу",
    "uploaded_at": "дата загрузки",
    "code": "код шаблона",
    "channel": "канал уведомления",
    "subject": "тема сообщения",
    "body_template": "шаблон текста",
    "recipient_id": "идентификатор получателя",
    "error_message": "сообщение об ошибке",
    "qr_code_data": "данные QR-кода",
    "confirmation_url": "ссылка для подтверждения",
    "external_id": "внешний идентификатор транзакции"
}

def get_col_desc(col_name):
    if col_name in column_desc_map:
        return column_desc_map[col_name]
    # Fallback: check if it contains common parts
    for key, val in column_desc_map.items():
        if key in col_name:
            return val
    return f"поле {col_name}"

# Detailed table descriptions from db_analysis.md plus extra advantages
table_descriptions = {
    "product_categories": "Справочник категорий товаров (розы, тюльпаны, букеты и т.д.). Позволяет структурировать каталог и упрощает поиск для клиентов.",
    "products": "Центральный каталог товаров — каждый цветок, букет, горшечное растение или аксессуар. Содержит актуальные цены и характеристики.",
    "users": "Аккаунты всех пользователей системы — клиентов, флористов, курьеров, администраторов. Обеспечивает безопасность и разграничение доступа.",
    "user_roles": "Реализация ролевой модели доступа (RBAC). Позволяет назначать пользователям одну или несколько ролей для гибкого управления правами.",
    "refresh_tokens": "Хранение токенов обновления для механизмов JWT-аутентификации. Повышает безопасность сессий без необходимости частого ввода пароля.",
    "stores": "Реестр филиалов цветочного магазина. Является корневой сущностью для большинства операционных процессов в системе.",
    "customers": "CRM-профили клиентов. Позволяет хранить историю взаимодействий, маркетинговые теги и персональные данные для улучшения сервиса.",
    "loyalty_accounts": "Счета программы лояльности клиентов. Отслеживает баланс баллов, уровень клиента и общую сумму трат.",
    "loyalty_transactions": "История всех операций по бонусным счетам (начисления, списания). Обеспечивает прозрачность программы лояльности.",
    "orders": "Центральная операционная сущность, описывающая заказы. Хранит информацию о статусе, оплате, способе доставки и привязке к клиенту.",
    "order_items": "Позиции заказов. Фиксирует состав заказа, цены и количество на момент совершения покупки (снимок данных).",
    "stock_balances": "Актуальные остатки товаров на конкретных складах и магазинах. Использует метод средневзвешенной себестоимости (WAC).",
    "stock_batches": "Партии поступивших товаров. Позволяет отслеживать сроки годности и автоматически списывать просроченную продукцию.",
    "stock_transactions": "Журнал всех движений товаров (приход, расход, списания). Является основой для проведения инвентаризаций и аудита.",
    "suppliers": "Справочник поставщиков продукции. Содержит контактные данные, платежные условия и рейтинг надежности.",
    "purchase_invoices": "Входящие приходные накладные. Служит основанием для оприходования товара на склад и создания новых партий.",
    "purchase_invoice_items": "Детализация приходных накладных по конкретным товарам с указанием заказанного и фактически принятого количества.",
    "employees": "Профили сотрудников организации. Содержит информацию о должности, дате приема и привязке к конкретному филиалу.",
    "employee_salary_configs": "Настройки алгоритмов расчета заработной платы. Поддерживает гибкие схемы: оклад, процент от продаж и бонусы.",
    "employee_timesheet": "Табель учета рабочего времени. Фиксирует время начала и окончания смены, а также общее количество отработанных часов.",
    "employee_salary_statements": "Итоговые ведомости на выплату зарплаты. Собирает данные из табеля и конфигураций для расчета финальных сумм.",
    "delivery_zones": "Географические зоны обслуживания с указанием стоимости доставки и минимальной суммы заказа для каждой зоны.",
    "delivery_slots": "Временные интервалы для планирования доставки. Позволяет балансировать нагрузку на курьерскую службу.",
    "delivery_tasks": "Задания на доставку конкретных заказов. Содержит статус выполнения, координаты и время прибытия курьера.",
    "financial_transactions": "Универсальный журнал финансовых операций. Регистрирует выручку, возвраты и операционные расходы предприятия.",
    "analytics_order_facts": "Агрегированные данные о завершенных заказах для быстрой аналитики выручки и маржинальности.",
    "analytics_cost_facts": "Денормализованные данные о затратах (закупки, зарплаты, списания) для анализа структуры расходов.",
    "media_files": "Метаданные загруженных медиа-файлов (фотографии товаров). Сами файлы хранятся в S3-совместимом хранилище.",
    "notification_templates": "Шаблоны уведомлений для различных каналов связи (SMS, Email, Push). Поддерживает переменную подстановку данных.",
    "notification_logs": "Журнал отправленных уведомлений. Позволяет отслеживать статус доставки сообщений и ошибки отправки.",
    "recipes": "Технологические карты (состав букетов). Позволяет автоматически списывать ингредиенты при продаже готовых композиций.",
    "recipe_items": "Конкретные составляющие рецепта с указанием необходимого количества каждого ингредиента.",
    "payments": "Транзакции платежных систем. Хранит данные об интеграции с внешними шлюзами и статусы онлайн-оплаты.",
    "shedlock": "Таблица распределенных блокировок. Гарантирует, что фоновые задачи выполняются только на одном экземпляре сервиса."
}

def generate_table_text(table_name, index):
    title = table_titles.get(table_name, table_name)
    desc = table_descriptions.get(table_name, "Описание таблицы.")
    cols = schema.get(table_name, [])
    
    attr_list = []
    for col in cols:
        cname = col['name']
        cdesc = get_col_desc(cname)
        attr_list.append(f"{cname} – {cdesc}")
    
    attr_str = "; ".join(attr_list)
    if attr_str:
        attr_str = f" Содержит следующие атрибуты: {attr_str}."
    
    return f"{index}. {title} ({table_name}). {desc}{attr_str}"

# Open doc
doc = Document(MY_DOC_PATH)

# Find insertion point
# We look for section "3.1 Описание таблиц"
target_p_idx = -1
for i, p in enumerate(doc.paragraphs):
    if "3.1" in p.text and "Описание" in p.text:
        target_p_idx = i
        break

if target_p_idx == -1:
    print("Could not find section 3.1. Inserting at the end.")
    target_p_idx = len(doc.paragraphs) - 1

# Tables to process in order
table_order = [
    "product_categories", "products", "users", "user_roles", "refresh_tokens",
    "stores", "customers", "loyalty_accounts", "loyalty_transactions",
    "orders", "order_items", "stock_balances", "stock_batches", "stock_transactions",
    "suppliers", "purchase_invoices", "purchase_invoice_items", "employees",
    "employee_salary_configs", "employee_timesheet", "employee_salary_statements",
    "delivery_zones", "delivery_slots", "delivery_tasks", "financial_transactions",
    "analytics_order_facts", "analytics_cost_facts", "media_files",
    "notification_templates", "notification_logs", "recipes", "recipe_items",
    "payments", "shedlock"
]

# Insert after the found paragraph
current_idx = target_p_idx + 1

for idx, tname in enumerate(table_order, 1):
    text = generate_table_text(tname, idx)
    
    # Create new paragraph
    new_p = doc.paragraphs[current_idx].insert_paragraph_before(text)
    
    # Apply formatting
    new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = new_p.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    
    # Apply font to runs
    for run in new_p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
    
    # If python-docx doesn't set font correctly due to theme, we might need a workaround
    # But usually this works.

# Save
doc.save(OUTPUT_PATH)
print(f"Generated {len(table_order)} table descriptions in {OUTPUT_PATH}")

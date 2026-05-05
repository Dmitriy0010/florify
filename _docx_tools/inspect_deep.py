"""
Deep inspection: read styles from XML, find insertion point in my doc,
and understand the exact formatting needed.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from lxml import etree
import re

EXAMPLE_PATH = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\5-текст отчета от Введения до Заключ.docx"
MY_DOC_PATH  = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\5-текст.docx"

def emu_to_cm(emu):
    if emu is None:
        return None
    return round(emu / 360000, 2)

def get_style_xml(doc, style_name):
    """Get the raw XML of a named style."""
    for s in doc.styles:
        if s.name == style_name:
            return etree.tostring(s.element, pretty_print=True).decode('utf-8')
    return None

def inspect_styles(path, label):
    doc = Document(path)
    print(f"\n{'='*60}")
    print(f"  STYLES in: {label}")
    print(f"{'='*60}")
    
    # Print key styles
    for style_name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']:
        xml = get_style_xml(doc, style_name)
        if xml:
            print(f"\n--- Style: {style_name} ---")
            # Extract font info
            fonts = re.findall(r'w:(?:ascii|hAnsi|cs)="([^"]+)"', xml)
            sizes = re.findall(r'<w:sz w:val="(\d+)"', xml)
            print(f"  Fonts found: {list(set(fonts))}")
            print(f"  Sizes (half-pts): {sizes}")
            # indent
            indents = re.findall(r'<w:ind ([^/]+)/>', xml)
            print(f"  Indents: {indents}")
            # spacing
            spacings = re.findall(r'<w:spacing ([^/]+)/>', xml)
            print(f"  Spacings: {spacings}")
    
    # Also list all custom styles
    print(f"\n--- All style names ---")
    for s in doc.styles:
        if s.type.name in ['PARAGRAPH']:
            print(f"  '{s.name}'")

def find_insertion_point(path):
    """Find where in my doc the table section should be inserted."""
    doc = Document(path)
    print(f"\n{'='*60}")
    print(f"  MY DOC - All paragraphs (to find insertion point)")
    print(f"  Total: {len(doc.paragraphs)}")
    print(f"{'='*60}")
    
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt:
            print(f"[{i:3d}] style='{p.style.name}' | '{txt[:90]}'")

def inspect_example_tables(path):
    """Find the table description section in example doc."""
    doc = Document(path)
    print(f"\n{'='*60}")
    print(f"  EXAMPLE - Looking for table description pattern")
    print(f"{'='*60}")
    
    in_section = False
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        # Look for anything that looks like table descriptions
        if any(w in txt for w in ['Photos', 'таблица', 'Таблица', 'хранен', 'карнав', 'масках', 'маскар', 'Универсальная']):
            in_section = True
        
        if in_section and txt:
            style = p.style.name
            pf = p.paragraph_format
            
            # Get font from runs or style
            font_name = None
            font_size = None
            bold = None
            for r in p.runs:
                if r.font.name:
                    font_name = r.font.name
                if r.font.size:
                    font_size = r.font.size
                if r.bold is not None:
                    bold = r.bold
                break
            
            print(f"[{i:3d}] style='{style}' bold={bold} font='{font_name}' size={font_size}")
            print(f"       sp_bef={emu_to_cm(pf.space_before)}cm sp_aft={emu_to_cm(pf.space_after)}cm fi={emu_to_cm(pf.first_line_indent)}cm li={emu_to_cm(pf.left_indent)}cm")
            print(f"       '{txt[:100]}'")
            print()
        
        if in_section and i > 300:
            break

inspect_styles(EXAMPLE_PATH, "EXAMPLE (Carnival)")
inspect_styles(MY_DOC_PATH, "MY DOC (Florify)")
find_insertion_point(MY_DOC_PATH)

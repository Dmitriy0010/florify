"""
Inspect the example docx (carnival costumes) to understand paragraph styles,
fonts, and formatting that we need to replicate.
"""
from docx import Document
from docx.shared import Pt
import json

EXAMPLE_PATH = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\5-текст отчета от Введения до Заключ.docx"
MY_DOC_PATH  = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\5-текст.docx"

def inspect_doc(path, label, max_paras=120):
    doc = Document(path)
    print(f"\n{'='*60}")
    print(f"  {label}")
    print(f"  {path}")
    print(f"  Total paragraphs: {len(doc.paragraphs)}")
    print(f"{'='*60}")

    for i, p in enumerate(doc.paragraphs[:max_paras]):
        style = p.style.name
        text_preview = p.text[:80].replace('\n', ' ')
        # Check first run formatting
        font_name = None
        font_size = None
        bold = None
        if p.runs:
            r = p.runs[0]
            font_name = r.font.name
            font_size = r.font.size
            bold = r.bold
        
        # Paragraph format
        pf = p.paragraph_format
        space_before = pf.space_before
        space_after = pf.space_after
        first_line = pf.first_line_indent
        left_indent = pf.left_indent
        alignment = pf.alignment

        print(f"[{i:3d}] style='{style}' | align={alignment} | sb={space_before} sa={space_after} fi={first_line} li={left_indent}")
        print(f"       font='{font_name}' size={font_size} bold={bold}")
        print(f"       text: '{text_preview}'")
        print()

def find_table_sections(path):
    """Find paragraphs that look like table descriptions in the example."""
    doc = Document(path)
    print(f"\n{'='*60}")
    print("  SEARCHING FOR TABLE DESCRIPTIONS IN EXAMPLE")
    print(f"{'='*60}")
    
    keywords = ["таблиц", "Таблиц", "хранен", "поле", "Поле", "первичн", "идентифик"]
    for i, p in enumerate(doc.paragraphs):
        if any(k in p.text for k in keywords):
            style = p.style.name
            pf = p.paragraph_format
            font_name = p.runs[0].font.name if p.runs else None
            font_size = p.runs[0].font.size if p.runs else None
            bold = p.runs[0].bold if p.runs else None
            print(f"[{i:3d}] style='{style}' font='{font_name}' size={font_size} bold={bold}")
            print(f"       text: '{p.text[:100]}'")

# Run inspection
inspect_doc(EXAMPLE_PATH, "EXAMPLE (Carnival)", max_paras=80)
find_table_sections(EXAMPLE_PATH)
inspect_doc(MY_DOC_PATH, "MY REPORT (Florify)", max_paras=60)

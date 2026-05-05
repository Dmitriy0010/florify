from docx import Document

doc = Document(r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\5-текст.docx")
for i in range(85, 100):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        print(f"[{i}] '{p.text}'")

"""
Extract XML of key styles from both documents and find exact 
paragraph index where to insert table descriptions in MY doc.
"""
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import re

EXAMPLE_PATH = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\5-текст отчета от Введения до Заключ.docx"
MY_DOC_PATH  = r"c:\Users\dda20\IdeaProjects\base-to-florify\florify\5-текст.docx"

# -------------------------------------------------------
# 1. Check styles XML from EXAMPLE
# -------------------------------------------------------
ex_doc = Document(EXAMPLE_PATH)
print("=== EXAMPLE: Normal style XML ===")
for s in ex_doc.styles:
    if s.name == 'Normal':
        xml = etree.tostring(s.element, pretty_print=True).decode('utf-8')
        print(xml[:3000])
        break

print("\n=== EXAMPLE: Heading 2 style XML ===")
for s in ex_doc.styles:
    if s.name == 'Heading 2':
        xml = etree.tostring(s.element, pretty_print=True).decode('utf-8')
        print(xml[:3000])
        break

# -------------------------------------------------------
# 2. Check MY DOC default paragraph format (from document XML)  
# -------------------------------------------------------
my_doc = Document(MY_DOC_PATH)
print("\n=== MY DOC: Normal style XML ===")
for s in my_doc.styles:
    if s.name == 'Normal':
        xml = etree.tostring(s.element, pretty_print=True).decode('utf-8')
        print(xml[:3000])
        break

print("\n=== MY DOC: Heading 2 style XML ===")
for s in my_doc.styles:
    if s.name == 'Heading 2':
        xml = etree.tostring(s.element, pretty_print=True).decode('utf-8')
        print(xml[:3000])
        break

# -------------------------------------------------------
# 3. Find exact insertion point in MY DOC
#    We look for "3.1" heading (section 3.1 of the report)
#    After that we need to find paragraphs before "3.2" 
# -------------------------------------------------------
print("\n=== MY DOC: paragraphs around section 3 ===")
paragraphs = my_doc.paragraphs
for i, p in enumerate(paragraphs):
    txt = p.text.strip()
    # Show paragraphs 85-115 (the DB section area)
    if 85 <= i <= 115:
        print(f"[{i:3d}] style='{p.style.name}' | '{txt[:80]}'")

# -------------------------------------------------------
# 4. Check what's in example around "Универсальная таблица" 
# -------------------------------------------------------
print("\n=== EXAMPLE: paragraphs around table descriptions ===")
ex_paragraphs = ex_doc.paragraphs
print(f"Total paragraphs in example: {len(ex_paragraphs)}")
for i, p in enumerate(ex_paragraphs):
    txt = p.text.strip()
    # Look for the start of table descriptions section
    if any(w in txt for w in ['Photos', 'Универсальная', 'карнавальн', 'маскарад', 'таблиц для']):
        # Print context: 5 before, 30 after
        start = max(0, i-2)
        end = min(len(ex_paragraphs), i+60)
        for j in range(start, end):
            pp = ex_paragraphs[j]
            pf = pp.paragraph_format
            def emu_to_cm(emu):
                if emu is None: return None
                return round(emu / 360000, 2)
            
            font_info = ''
            if pp.runs:
                r = pp.runs[0]
                font_info = f"font='{r.font.name}' sz={r.font.size} bold={r.bold} italic={r.italic}"
            
            print(f"[{j:3d}] style='{pp.style.name}' align={pf.alignment} sb={emu_to_cm(pf.space_before)} sa={emu_to_cm(pf.space_after)} fi={emu_to_cm(pf.first_line_indent)} li={emu_to_cm(pf.left_indent)}")
            print(f"       {font_info}")
            print(f"       '{pp.text[:100]}'")
        break
